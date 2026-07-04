import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt


REQUIRED_SHEETS = [
    "General",
    "Down Time",
    "Book Wise Details",
]


LPRS_RULES = {
    "TOIM_MP_1": "00:00",
    "MTM_MP_1": "23:15",
    "TOITH_MP_1": "00:30",
    "TOIVP_MP_1": "00:15",
}


# -------------------------------------------------
# BASIC HELPERS
# -------------------------------------------------
def find_column(df, possible_names):
    df_columns_clean = {
        str(col).strip().lower(): col
        for col in df.columns
    }

    for name in possible_names:
        clean_name = name.strip().lower()
        if clean_name in df_columns_clean:
            return df_columns_clean[clean_name]

    return None


def format_time(value):
    if pd.isna(value):
        return ""

    try:
        value = pd.to_datetime(value)
        return value.strftime("%H:%M hrs")
    except Exception:
        return str(value)


def clean_complexity(value):
    if pd.isna(value):
        return ""

    text = str(value).strip()

    if "-" in text:
        return text.split("-", 1)[1].strip()

    return text


def clean_reason_text(value):
    if pd.isna(value):
        return ""

    return str(value).strip()


# -------------------------------------------------
# TIME / DELAY CALCULATION
# -------------------------------------------------
def parse_time_to_minutes(value):
    if pd.isna(value):
        return None

    try:
        dt = pd.to_datetime(value)
        return dt.hour * 60 + dt.minute
    except Exception:
        pass

    try:
        text = str(value).strip().replace("hrs", "").strip()
        parts = text.split(":")
        hour = int(parts[0])
        minute = int(parts[1])
        return hour * 60 + minute
    except Exception:
        return None


def calculate_page_release_delay(product_name, lpr_value):
    product_text = str(product_name).strip()

    if product_text not in LPRS_RULES:
        return "", "", ""

    lprs_text = LPRS_RULES[product_text]

    lpr_minutes = parse_time_to_minutes(lpr_value)
    lprs_minutes = parse_time_to_minutes(lprs_text)

    if lpr_minutes is None or lprs_minutes is None:
        return format_time(lpr_value), f"{lprs_text} hrs", ""

    delay = lpr_minutes - lprs_minutes

    if delay < 0:
        delay += 24 * 60

    return format_time(lpr_value), f"{lprs_text} hrs", f"{delay} mins"


# -------------------------------------------------
# EXCLUSION RULE
# -------------------------------------------------
def should_exclude_downtime(row, related_col=None, reason_col=None):
    """
    Exclude Reflong / RMC / Editorial downtime.
    Checks both Related and Reason columns.
    """

    related_text = ""
    reason_text = ""

    if related_col is not None:
        related_text = str(row.get(related_col, "")).strip().lower()

    if reason_col is not None:
        reason_text = str(row.get(reason_col, "")).strip().lower()

    combined_text = related_text + " " + reason_text

    exclude_words = [
        "reflong",
        "rmc",
        "editorial",
    ]

    for word in exclude_words:
        if word in combined_text:
            return True

    return False


# -------------------------------------------------
# IDENTIFY LAST FINISHED MAIN EDITION
# -------------------------------------------------
def identify_last_finished_main_edition(general_df):
    main_supp_col = find_column(
        general_df,
        ["Main/Supplement", "Main Supplement", "Main_Supplement"]
    )

    production_end_col = find_column(
        general_df,
        ["Last Production End", "Production End"]
    )

    production_end_date_col = find_column(
        general_df,
        ["Production End Date", "Last Production End Date"]
    )

    product_code_col = find_column(
        general_df,
        ["Product Name", "Edition", "Products", "Product"]
    )

    machine_col = find_column(
        general_df,
        ["Machine", "Machine Name"]
    )

    runid_col = find_column(
        general_df,
        ["Runid", "Run ID", "RunId"]
    )

    required_columns = {
        "Main/Supplement": main_supp_col,
        "Production End": production_end_col,
        "Production End Date": production_end_date_col,
        "Product Name / Edition / Products": product_code_col,
        "Machine": machine_col,
        "Runid": runid_col,
    }

    missing_columns = [
        name for name, col in required_columns.items()
        if col is None
    ]

    if missing_columns:
        return None, missing_columns

    main_df = general_df[
        general_df[main_supp_col].astype(str).str.strip().str.lower() == "main"
    ].copy()

    if main_df.empty:
        return pd.DataFrame(), []

    main_df["_production_end_datetime"] = pd.to_datetime(
        main_df[production_end_date_col].astype(str).str.strip()
        + " "
        + main_df[production_end_col].astype(str).str.strip(),
        errors="coerce",
        dayfirst=True
    )

    main_df = main_df.dropna(subset=["_production_end_datetime"])

    if main_df.empty:
        return pd.DataFrame(), []

    latest_finish_datetime = main_df["_production_end_datetime"].max()

    last_finished_df = main_df[
        main_df["_production_end_datetime"] == latest_finish_datetime
    ].copy()

    return last_finished_df, []


# -------------------------------------------------
# BOOK WISE DETAILS MATCHING
# -------------------------------------------------
def get_bookwise_info(bookwise_df, runid):
    runid_col = find_column(
        bookwise_df,
        ["Runid", "Run ID", "RunId"]
    )

    edition_col = find_column(
        bookwise_df,
        ["Edition", "Edition Name"]
    )

    last_tiff_col = find_column(
        bookwise_df,
        ["Last Tiff", "Last Tiff Edition", "LPR"]
    )

    complexity_col = find_column(
        bookwise_df,
        ["Complexities", "Complexity"]
    )

    required_columns = {
        "Runid": runid_col,
        "Edition": edition_col,
        "Last Tiff": last_tiff_col,
        "Complexities": complexity_col,
    }

    missing_columns = [
        name for name, col in required_columns.items()
        if col is None
    ]

    if missing_columns:
        return None, missing_columns

    matched_df = bookwise_df[
        bookwise_df[runid_col].astype(str).str.strip() == str(runid).strip()
    ]

    if matched_df.empty:
        return None, []

    row = matched_df.iloc[0]

    return {
        "edition": row[edition_col],
        "last_tiff": row[last_tiff_col],
        "complexity": clean_complexity(row[complexity_col]),
    }, []


def get_issue_date(last_finished_df):
    production_end_date_col = find_column(
        last_finished_df,
        ["Production End Date", "Last Production End Date"]
    )

    if production_end_date_col is None:
        return ""

    try:
        latest_date = pd.to_datetime(
            last_finished_df[production_end_date_col].iloc[0],
            dayfirst=True
        )
        return latest_date.strftime("%d-%m-%Y")
    except Exception:
        return ""


# -------------------------------------------------
# VALID DOWNTIME CALCULATION FROM DOWN TIME SHEET
# -------------------------------------------------
def get_valid_downtime_from_downtime_sheet(downtime_df, runids=None):
    """
    Calculates valid downtime from Down Time sheet.
    Excludes Reflong / RMC / Editorial.
    """

    if runids is None:
        runids = []

    if not isinstance(runids, list):
        runids = [runids]

    runids = [str(x).strip() for x in runids]

    runid_col = find_column(
        downtime_df,
        ["Runid", "Run ID", "RunId"]
    )

    main_supp_col = find_column(
        downtime_df,
        ["Main/Supplement", "Main Supplement", "Main_Supplement"]
    )

    related_col = find_column(
        downtime_df,
        ["Related"]
    )

    reason_col = find_column(
        downtime_df,
        ["Reason", "Downtime Reason", "Down Time Reason"]
    )

    downtime_col = find_column(
        downtime_df,
        ["Total Downtime", "Total DownTime", "Downtime", "Down Time"]
    )

    required_columns = {
        "Runid": runid_col,
        "Main/Supplement": main_supp_col,
        "Down Time Minutes": downtime_col,
    }

    missing_columns = [
        name for name, col in required_columns.items()
        if col is None
    ]

    if missing_columns:
        return 0, missing_columns

    df = downtime_df[
        downtime_df[main_supp_col].astype(str).str.strip().str.lower() == "main"
    ].copy()

    if runids:
        df = df[
            df[runid_col].astype(str).str.strip().isin(runids)
        ].copy()

    if df.empty:
        return 0, []

    df = df[
        ~df.apply(
            lambda row: should_exclude_downtime(row, related_col, reason_col),
            axis=1
        )
    ].copy()

    if df.empty:
        return 0, []

    df[downtime_col] = pd.to_numeric(
        df[downtime_col],
        errors="coerce"
    ).fillna(0)

    total_downtime = int(round(df[downtime_col].sum()))

    return total_downtime, []


def calculate_machine_wise_downtime(general_df, downtime_df):
    """
    Calculates machine-wise downtime for Main editions.
    Machine and Runid list comes from General sheet.
    Downtime comes from Down Time sheet after excluding Reflong/RMC/Editorial.
    """

    main_supp_col = find_column(
        general_df,
        ["Main/Supplement", "Main Supplement", "Main_Supplement"]
    )

    machine_col = find_column(
        general_df,
        ["Machine", "Machine Name"]
    )

    runid_col = find_column(
        general_df,
        ["Runid", "Run ID", "RunId"]
    )

    required_columns = {
        "Main/Supplement": main_supp_col,
        "Machine": machine_col,
        "Runid": runid_col,
    }

    missing_columns = [
        name for name, col in required_columns.items()
        if col is None
    ]

    if missing_columns:
        return None, missing_columns

    main_general_df = general_df[
        general_df[main_supp_col].astype(str).str.strip().str.lower() == "main"
    ].copy()

    if main_general_df.empty:
        return [], []

    machine_list = (
        main_general_df[machine_col]
        .dropna()
        .astype(str)
        .str.strip()
        .unique()
        .tolist()
    )

    result = []

    for machine_name in machine_list:
        machine_runids = (
            main_general_df[
                main_general_df[machine_col].astype(str).str.strip() == machine_name
            ][runid_col]
            .dropna()
            .astype(str)
            .str.strip()
            .unique()
            .tolist()
        )

        valid_downtime, missing_dt_columns = get_valid_downtime_from_downtime_sheet(
            downtime_df,
            runids=machine_runids
        )

        if missing_dt_columns:
            return None, missing_dt_columns

        result.append({
            "machine": machine_name,
            "downtime": valid_downtime
        })

    return result, []


# -------------------------------------------------
# DIRECTOR-LEVEL DELAY REASON WRITING
# -------------------------------------------------
def normalize_delay_reason(raw_reason):
    text = clean_reason_text(raw_reason).lower()

    if text == "":
        return ""

    if "web break" in text or "webbreak" in text:
        if "blanket" in text or "delamination" in text:
            return "web break caused by blanket delamination"
        return "web break"

    if "blanket" in text or "delamination" in text:
        return "web break caused by blanket delamination"

    if "safe stop" in text:
        return "safe stop tripping"

    if "folder jam" in text:
        return "folder jam"

    if "paper jam" in text:
        return "folder jam"

    if "paper drifting" in text or "paper drift" in text:
        return "paper drifting"

    if "production stop" in text:
        return "production stop"

    if "electrical" in text:
        return "electrical stoppage"

    if "plate" in text:
        return "plate delay"

    if "tiff" in text or "page" in text:
        return "page release delay"

    return clean_reason_text(raw_reason)


def get_reason_breakup(reason_df, reason_col, downtime_col):
    if reason_df.empty:
        return []

    temp_df = reason_df.copy()

    temp_df[downtime_col] = pd.to_numeric(
        temp_df[downtime_col],
        errors="coerce"
    ).fillna(0)

    temp_df["_reason_clean"] = temp_df[reason_col].apply(normalize_delay_reason)

    temp_df = temp_df[temp_df["_reason_clean"] != ""]

    if temp_df.empty:
        return []

    reason_summary = (
        temp_df
        .groupby("_reason_clean")[downtime_col]
        .sum()
        .reset_index()
        .sort_values(by=downtime_col, ascending=False)
    )

    result = []

    for _, row in reason_summary.iterrows():
        mins = int(round(row[downtime_col]))

        if mins > 0:
            result.append({
                "reason": row["_reason_clean"],
                "mins": mins
            })

    return result


def format_breakup_with_minutes(breakup):
    if not breakup:
        return ""

    parts = []

    for item in breakup:
        minute_word = "min" if item["mins"] == 1 else "mins"
        parts.append(f"{item['mins']} {minute_word} due to {item['reason']}")

    if len(parts) == 1:
        return parts[0]

    if len(parts) == 2:
        return " and ".join(parts)

    return ", ".join(parts[:-1]) + ", and " + parts[-1]


def format_reason_names_only(breakup):
    if not breakup:
        return ""

    names = []

    for item in breakup:
        reason = item["reason"]

        if reason == "folder jam":
            reason = "folder jams"
        elif reason == "web break":
            reason = "web breaks"

        names.append(reason)

    names = list(dict.fromkeys(names))

    if len(names) == 1:
        return names[0]

    if len(names) == 2:
        return " and ".join(names)

    return ", ".join(names[:-1]) + ", and " + names[-1]


def make_machine_text(machine_list):
    if not machine_list:
        return ""

    if len(machine_list) == 1:
        return machine_list[0]

    if len(machine_list) == 2:
        return f"{machine_list[0]} and {machine_list[1]}"

    return ", ".join(machine_list)


def join_sentence_parts(parts):
    if not parts:
        return ""

    if len(parts) == 1:
        return parts[0]

    if len(parts) == 2:
        return " and ".join(parts)

    return ", ".join(parts[:-1]) + ", and " + parts[-1]


def generate_delayed_finish_reason(general_df, downtime_df, last_finished_df):
    """
    Generates non-generic, human-readable boss-facing delay reason.
    Intro sentence is built from actual downtime reasons.
    """

    general_runid_col = find_column(
        general_df,
        ["Runid", "Run ID", "RunId"]
    )

    general_machine_col = find_column(
        general_df,
        ["Machine", "Machine Name"]
    )

    general_product_col = find_column(
        general_df,
        ["Products", "Product Name", "Product"]
    )

    print_delay_reason_col = find_column(
        general_df,
        ["Print finish delay Reason", "Print Finish Delay Reason", "Delay Reason"]
    )

    downtime_runid_col = find_column(
        downtime_df,
        ["Runid", "Run ID", "RunId"]
    )

    folder_col = find_column(
        downtime_df,
        ["Folder"]
    )

    downtime_main_supp_col = find_column(
        downtime_df,
        ["Main/Supplement", "Main Supplement", "Main_Supplement"]
    )

    related_col = find_column(
        downtime_df,
        ["Related"]
    )

    downtime_reason_col = find_column(
        downtime_df,
        ["Reason", "Downtime Reason", "Down Time Reason"]
    )

    downtime_minutes_col = find_column(
        downtime_df,
        ["Total Downtime", "Total DownTime", "Downtime", "Down Time"]
    )

    required_columns = {
        "General Runid": general_runid_col,
        "General Machine": general_machine_col,
        "General Products": general_product_col,
        "Down Time Runid": downtime_runid_col,
        "Folder": folder_col,
        "Down Time Main/Supplement": downtime_main_supp_col,
        "Down Time Reason": downtime_reason_col,
        "Down Time Minutes": downtime_minutes_col,
    }

    missing_columns = [
        name for name, col in required_columns.items()
        if col is None
    ]

    if missing_columns:
        return "", missing_columns

    last_runids = [
        str(value).strip()
        for value in last_finished_df[general_runid_col].tolist()
    ]

    last_machines = [
        str(value).strip()
        for value in last_finished_df[general_machine_col].dropna().unique().tolist()
    ]

    last_products = [
        str(value).strip()
        for value in last_finished_df[general_product_col].dropna().unique().tolist()
    ]

    if last_products:
        edition_name_for_reason = last_products[0]
    else:
        edition_name_for_reason = "the last edition"

    machine_text = make_machine_text(last_machines)

    downtime_main_df = downtime_df[
        downtime_df[downtime_main_supp_col].astype(str).str.strip().str.lower() == "main"
    ].copy()

    if downtime_main_df.empty:
        return (
            f"Sir, the late finish of {edition_name_for_reason} on {machine_text} was recorded without any valid downtime entry.",
            []
        )

    downtime_main_df = downtime_main_df[
        ~downtime_main_df.apply(
            lambda row: should_exclude_downtime(row, related_col, downtime_reason_col),
            axis=1
        )
    ].copy()

    direct_downtime_df = downtime_main_df[
        downtime_main_df[downtime_runid_col].astype(str).str.strip().isin(last_runids)
    ].copy()

    possible_folders_df = downtime_df[
        downtime_df[downtime_runid_col].astype(str).str.strip().isin(last_runids)
    ]

    folders = (
        possible_folders_df[folder_col]
        .dropna()
        .astype(str)
        .str.strip()
        .unique()
        .tolist()
    )

    same_folder_df = downtime_main_df[
        downtime_main_df[folder_col].astype(str).str.strip().isin(folders)
    ].copy()

    cascading_df = same_folder_df[
        ~same_folder_df[downtime_runid_col].astype(str).str.strip().isin(last_runids)
    ].copy()

    reason_lines = []
    direct_machine_sentences = []
    machine_reason_parts = []

    for _, last_row in last_finished_df.iterrows():
        runid = str(last_row[general_runid_col]).strip()
        machine = str(last_row[general_machine_col]).strip()

        machine_direct_df = direct_downtime_df[
            direct_downtime_df[downtime_runid_col].astype(str).str.strip() == runid
        ].copy()

        machine_breakup = get_reason_breakup(
            machine_direct_df,
            downtime_reason_col,
            downtime_minutes_col
        )

        machine_total = sum(item["mins"] for item in machine_breakup)

        if machine_breakup and machine_total > 0:
            reason_names = format_reason_names_only(machine_breakup)
            machine_reason_parts.append(f"{reason_names} on {machine}")

            direct_machine_sentences.append(
                f"On {machine}, the edition recorded {machine_total} mins downtime, including "
                f"{format_breakup_with_minutes(machine_breakup)}."
            )
        else:
            direct_machine_sentences.append(
                f"On {machine}, no valid direct downtime was recorded after excluding Reflong/RMC/Editorial downtime."
            )

    if machine_reason_parts:
        main_reason_text = join_sentence_parts(machine_reason_parts)

        reason_lines.append(
            f"Sir, the late finish of {edition_name_for_reason} on {machine_text} "
            f"was mainly due to {main_reason_text}."
        )
    else:
        reason_lines.append(
            f"Sir, the late finish of {edition_name_for_reason} on {machine_text} "
            f"was not linked to any valid direct downtime after excluding Reflong/RMC/Editorial entries."
        )

    if direct_machine_sentences:
        reason_lines.extend(direct_machine_sentences)

    cascading_breakup = get_reason_breakup(
        cascading_df,
        downtime_reason_col,
        downtime_minutes_col
    )

    if cascading_breakup:
        cascading_reason_names = format_reason_names_only(cascading_breakup)
        reason_lines.append(
            f"Earlier valid stoppages on the same folder, mainly due to {cascading_reason_names}, "
            f"also added pressure on the final run."
        )

    additional_reasons = []

    if print_delay_reason_col is not None:
        for _, row in last_finished_df.iterrows():
            reason_text = clean_reason_text(row.get(print_delay_reason_col, ""))
            if reason_text and reason_text.lower() != "nan":
                additional_reasons.append(reason_text)

    additional_reasons = list(dict.fromkeys(additional_reasons))

    if additional_reasons:
        reason_lines.append(
            "Additional delay reason: " + "; ".join(additional_reasons) + "."
        )

    final_reason = " ".join(reason_lines)

    return final_reason, []


# -------------------------------------------------
# WORD FILE CREATION
# -------------------------------------------------
def create_word_report(report_text):
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = report_text.split("\n")

    for line in lines:
        clean_line = line.strip()

        if clean_line == "":
            document.add_paragraph("")
            continue

        paragraph = document.add_paragraph()

        if clean_line.startswith("Airoli Plant Issue Dated:"):
            run = paragraph.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(13)

        elif clean_line in [
            "Page Release Delay:",
            "Machine-wise Total Downtime for Main Editions:",
            "Reasons for delayed finish:",
        ]:
            run = paragraph.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(11)

        else:
            run = paragraph.add_run(clean_line)
            run.font.size = Pt(11)

    word_buffer = BytesIO()
    document.save(word_buffer)
    word_buffer.seek(0)

    return word_buffer


# -------------------------------------------------
# MAIN STREAMLIT MODULE
# -------------------------------------------------
def show_pf_delay_report():
    st.markdown("## PF Delay Report")
    st.caption("Generate Director-level Print Finished delay report from Production Excel file.")

    uploaded_file = st.file_uploader(
        "Upload Production Report Excel File",
        type=["xlsx", "xls"],
        key="pf_delay_report_uploader"
    )

    if uploaded_file is None:
        st.info("Upload the daily production Excel file to generate PF Delay Report.")
        return

    try:
        excel_file = pd.ExcelFile(uploaded_file)
        available_sheets = excel_file.sheet_names

        missing_sheets = [
            sheet for sheet in REQUIRED_SHEETS
            if sheet not in available_sheets
        ]

        if missing_sheets:
            st.error("Required sheet missing in uploaded file.")
            st.write("Missing sheet(s):")
            st.write(missing_sheets)
            return

        general_df = pd.read_excel(uploaded_file, sheet_name="General")
        downtime_df = pd.read_excel(uploaded_file, sheet_name="Down Time")
        bookwise_df = pd.read_excel(uploaded_file, sheet_name="Book Wise Details")

        st.success("File uploaded successfully.")

        last_finished_df, missing_columns = identify_last_finished_main_edition(general_df)

        if missing_columns:
            st.error("Required column missing in General sheet.")
            st.write("Missing column(s):")
            st.write(missing_columns)
            return

        if last_finished_df is None or last_finished_df.empty:
            st.warning("No Main edition records found with valid print finish time.")
            return

        product_code_col = find_column(
            last_finished_df,
            ["Product Name", "Edition", "Products", "Product"]
        )

        machine_col = find_column(
            last_finished_df,
            ["Machine", "Machine Name"]
        )

        production_end_col = find_column(
            last_finished_df,
            ["Last Production End", "Production End"]
        )

        runid_col = find_column(
            last_finished_df,
            ["Runid", "Run ID", "RunId"]
        )

        issue_date = get_issue_date(last_finished_df)

        report_lines = []
        report_lines.append(f"Airoli Plant Issue Dated: {issue_date}")
        report_lines.append("")

        first_lpr = ""
        first_lprs = ""
        first_delay = ""

        for _, row in last_finished_df.iterrows():
            runid = row[runid_col]
            product_code = row[product_code_col]

            book_info, book_missing_columns = get_bookwise_info(bookwise_df, runid)

            if book_missing_columns:
                st.error("Required column missing in Book Wise Details sheet.")
                st.write("Missing column(s):")
                st.write(book_missing_columns)
                return

            if book_info is None:
                edition_name = str(row[product_code_col])
                complexity = "Bookwise Details missing"
                lpr = ""
                lprs = ""
                delay = ""
            else:
                edition_name = book_info["edition"]
                complexity = book_info["complexity"]
                lpr, lprs, delay = calculate_page_release_delay(
                    product_code,
                    book_info["last_tiff"]
                )

            if first_lpr == "":
                first_lpr = lpr
                first_lprs = lprs
                first_delay = delay

            valid_last_downtime, last_dt_missing_columns = get_valid_downtime_from_downtime_sheet(
                downtime_df,
                runids=[runid]
            )

            if last_dt_missing_columns:
                st.error("Required column missing for last edition downtime calculation.")
                st.write("Missing column(s):")
                st.write(last_dt_missing_columns)
                return

            report_lines.append(f"Last Edition: {edition_name}")
            report_lines.append(f"Print Finish: {format_time(row[production_end_col])}")
            report_lines.append(f"Machine: {row[machine_col]}")
            report_lines.append(f"Total Downtime: {valid_last_downtime} mins")
            report_lines.append(f"Complexity: {complexity}")
            report_lines.append("")

        report_lines.append("Page Release Delay:")
        report_lines.append(f"LPR: {first_lpr}")
        report_lines.append(f"LPRS: {first_lprs}")
        report_lines.append(f"Delay: {first_delay}")
        report_lines.append("")

        machine_wise_downtime, machine_missing_columns = calculate_machine_wise_downtime(
            general_df,
            downtime_df
        )

        if machine_missing_columns:
            st.error("Required column missing for machine-wise downtime calculation.")
            st.write("Missing column(s):")
            st.write(machine_missing_columns)
            return

        report_lines.append("Machine-wise Total Downtime for Main Editions:")

        for item in machine_wise_downtime:
            report_lines.append(f"{item['machine']}: {item['downtime']} mins")

        report_lines.append("")

        delayed_reason, reason_missing_columns = generate_delayed_finish_reason(
            general_df,
            downtime_df,
            last_finished_df
        )

        if reason_missing_columns:
            st.error("Required column missing for delayed finish reason.")
            st.write("Missing column(s):")
            st.write(reason_missing_columns)
            return

        report_lines.append("Reasons for delayed finish:")
        report_lines.append(delayed_reason)
        report_lines.append("")
        report_lines.append("Copies delivered after 04:00 am:")

        report_text = "\n".join(report_lines)

        st.markdown("### PF Delay Report Preview")

        st.text_area(
            "Generated Report Text",
            value=report_text,
            height=650,
            key="pf_delay_report_text_preview"
        )

        word_file = create_word_report(report_text)

        download_file_name = f"PF_Delay_Report_{issue_date}.docx"

        st.download_button(
            label="Download Word Report",
            data=word_file,
            file_name=download_file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="pf_delay_report_word_download"
        )

    except Exception as e:
        st.error("Unable to generate PF Delay Report.")
        st.exception(e)
